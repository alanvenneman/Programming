import docx
import os


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


document = docx.Document()
p = document.add_paragraph('A plain paragraph having some ')
in_geodatabase = "S123_dd48327d928e451fa4a0b063cfc90e61_FGDB.zip"  # arcpy.GetParameterAsText(0)
geodatabase_path = os.path.join(r"\\vs512\e$\Scripts\BatchPhotoTesting", in_geodatabase)
gdb_split = in_geodatabase.split('_')
itemID = gdb_split[1]
url_string = r"https://survey123.arcgis.com/surveys/{}/data?report=format:docx".format(itemID)
add_hyperlink(p, 'Link to my site', url_string)
document.save('demo_hyperlink.docx')